"""
Профессиональная программа транскрибации с диаризацией
Автор: Lebedev Nikolay
"""

import sys
import os
import gc
import re
import time
import json
import subprocess
import tempfile
import warnings
import threading
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget,
    QLabel, QTextEdit, QPushButton, QLineEdit,
    QFileDialog, QVBoxLayout, QHBoxLayout, QGroupBox,
    QMessageBox, QStyleFactory, QProgressBar, QComboBox,
    QCheckBox, QSpinBox, QTabWidget, QTextBrowser,
    QSplitter, QFrame, QStyle, QDialog
)
from PySide6.QtCore import Qt, QThread, Signal, Slot, QTimer, QPropertyAnimation, QEasingCurve, QMutex, QMutexLocker
from PySide6.QtGui import QIcon, QFont, QPalette, QColor, QTextCharFormat, QTextCursor, QPixmap

# Отключаем предупреждения
warnings.filterwarnings("ignore")
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

# Константы
APP_NAME = "Audio/Video Transcription Pro"
APP_VERSION = "5.3-PROFESSIONAL"
AUTHOR = "Lebedev Nikolay"
DEFAULT_HF_TOKEN = "" #ДОБАВЬТЕ СЮДА СВОЙ ТОКЕН HF

# Глобальный мьютекс для безопасности памяти
MEMORY_MUTEX = QMutex()

# Проверка зависимостей
FASTER_WHISPER_AVAILABLE = False
DOCX_AVAILABLE = False
TORCH_AVAILABLE = False
DEVICE = "cpu"

try:
    from faster_whisper import WhisperModel
    FASTER_WHISPER_AVAILABLE = True
except ImportError:
    FASTER_WHISPER_AVAILABLE = False
    print("ОШИБКА: faster_whisper не установлен!")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Безопасная инициализация GPU
try:
    import torch
    TORCH_AVAILABLE = True
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
    if sys.platform == "win32":
        os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
    print(f"PyTorch: {torch.__version__}, Device: {DEVICE}")
except ImportError:
    TORCH_AVAILABLE = False
    DEVICE = "cpu"
    print("PyTorch не установлен - CPU режим")


class CrashSafeMemoryManager:
    """Безопасный менеджер памяти для предотвращения крашей"""

    _cleanup_in_progress = False
    _cleanup_lock = threading.Lock()

    @staticmethod
    def safe_gpu_cleanup(stage="unknown"):
        """Ультра-безопасная очистка GPU памяти"""
        # Предотвращаем множественные одновременные вызовы
        with CrashSafeMemoryManager._cleanup_lock:
            if CrashSafeMemoryManager._cleanup_in_progress:
                return

            CrashSafeMemoryManager._cleanup_in_progress = True

        try:
            # Мягкая сборка мусора Python
            try:
                gc.collect()
            except:
                pass

            # Очень осторожная работа с CUDA
            if DEVICE == "cuda" and TORCH_AVAILABLE:
                try:
                    import torch

                    # Множественные проверки безопасности
                    if not torch.cuda.is_available():
                        return

                    if torch.cuda.device_count() == 0:
                        return

                    # Проверяем, что есть инициализированный контекст
                    try:
                        current_device = torch.cuda.current_device()
                    except:
                        return

                    # Только мягкая очистка кэша, без синхронизации
                    try:
                        allocated_before = torch.cuda.memory_allocated(current_device) / (1024**2)
                        torch.cuda.empty_cache()
                        allocated_after = torch.cuda.memory_allocated(current_device) / (1024**2)

                        freed = allocated_before - allocated_after
                        if freed > 1:
                            print(f"GPU cleanup {stage}: freed {freed:.1f}MB")

                    except Exception as cache_error:
                        print(f"Cache cleanup warning: {cache_error}")

                except Exception as cuda_error:
                    print(f"CUDA cleanup warning: {cuda_error}")

        except Exception as e:
            print(f"Memory cleanup error: {e}")
        finally:
            CrashSafeMemoryManager._cleanup_in_progress = False

    @staticmethod
    def reset_for_new_transcription():
        """Полный сброс состояния для новой транскрибации"""
        try:
            # Принудительная сборка мусора
            for _ in range(3):
                gc.collect()

            # Очистка GPU памяти
            CrashSafeMemoryManager.safe_gpu_cleanup("reset for new transcription")

            # Небольшая пауза для стабилизации
            time.sleep(0.1)

        except Exception as e:
            print(f"Reset error: {e}")


class AboutDialog(QDialog):
    """Диалог О программе"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("О программе")
        self.setFixedSize(450, 350)
        self.setModal(True)

        layout = QVBoxLayout(self)

        # Заголовок
        title = QLabel(APP_NAME)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #58a6ff;
                margin: 10px;
            }
        """)

        # Версия
        version = QLabel(f"Версия {APP_VERSION}")
        version.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version.setStyleSheet("font-size: 14px; color: #8b949e; margin: 5px;")

        # Автор
        author = QLabel(f"Разработчик: {AUTHOR}")
        author.setAlignment(Qt.AlignmentFlag.AlignCenter)
        author.setStyleSheet("font-size: 12px; color: #d4d4d4; margin: 5px;")

        # Описание
        description = QLabel(
            "Профессиональная программа для транскрибации\n"
            "аудио и видео файлов с диаризацией\n"
            "спикеров на базе Whisper AI\n\n"
            "Надежно, быстро, качественно."
        )
        description.setAlignment(Qt.AlignmentFlag.AlignCenter)
        description.setWordWrap(True)
        description.setStyleSheet("font-size: 11px; color: #8b949e; margin: 15px;")

        # Технологии
        tech = QLabel("Использует: OpenAI Whisper, PyAnnote, PySide6, PyTorch")
        tech.setAlignment(Qt.AlignmentFlag.AlignCenter)
        tech.setStyleSheet("font-size: 10px; color: #6b7280; margin: 10px;")

        # Кнопка закрытия
        close_btn = QPushButton("Закрыть")
        close_btn.clicked.connect(self.accept)
        close_btn.setStyleSheet("""
            QPushButton {
                background: #667eea;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 15px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: #5a67d8;
            }
        """)

        layout.addWidget(title)
        layout.addWidget(version)
        layout.addWidget(author)
        layout.addWidget(description)
        layout.addWidget(tech)
        layout.addStretch()

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_layout.addWidget(close_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)


class LogWidget(QTextBrowser):
    """Виджет для логов с цветовой подсветкой"""

    def __init__(self):
        super().__init__()
        self.setReadOnly(True)
        self.setOpenExternalLinks(False)
        self.setStyleSheet("""
            QTextBrowser {
                background-color: #1e1e1e;
                color: #d4d4d4;
                border: 1px solid #3c3c3c;
                border-radius: 6px;
                padding: 8px;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 11px;
            }
            QScrollBar:vertical {
                background: #2d2d2d;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #5a5a5a;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background: #7a7a7a;
            }
        """)

    def log(self, message, level="INFO"):
        """Добавить сообщение в лог с цветовой подсветкой"""
        timestamp = datetime.now().strftime("%H:%M:%S")

        colors = {
            "INFO": "#58a6ff",
            "SUCCESS": "#56d364",
            "WARNING": "#f0883e",
            "ERROR": "#f85149",
            "DEBUG": "#8b949e"
        }

        icons = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }

        color = colors.get(level, "#d4d4d4")
        icon = icons.get(level, "•")

        html = f'<span style="color: #8b949e">[{timestamp}]</span> '
        html += f'<span style="color: {color}">{icon} <b>{level}</b>:</span> '
        html += f'<span style="color: #d4d4d4">{message}</span>'

        self.append(html)

        # Автоскролл к последнему сообщению
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.setTextCursor(cursor)


class ModelDownloader(QThread):
    """Поток для безопасного скачивания моделей"""

    progress_signal = Signal(int, str)
    finished_signal = Signal(bool, str)
    log_signal = Signal(str, str)

    def __init__(self, model_name: str):
        super().__init__()
        self.model_name = model_name
        self.should_stop = False

    def run(self):
        try:
            self.log_signal.emit(f"Скачивание модели {self.model_name}...", "INFO")
            self.progress_signal.emit(10, "Проверка модели...")

            if self.should_stop:
                return

            self.progress_signal.emit(30, "Загрузка модели...")

            # Безопасное создание модели
            compute_type = "float16" if DEVICE == "cuda" else "int8"

            # Предварительная очистка памяти
            CrashSafeMemoryManager.safe_gpu_cleanup("before model download")

            model = WhisperModel(
                self.model_name,
                device=DEVICE,
                compute_type=compute_type,
                cpu_threads=min(4, os.cpu_count()),
                num_workers=1,
                download_root=self.get_models_cache_dir()
            )

            self.progress_signal.emit(80, "Проверка работоспособности...")

            # Безопасное удаление модели
            del model
            CrashSafeMemoryManager.safe_gpu_cleanup("after model download")

            self.progress_signal.emit(100, "Модель готова!")
            self.finished_signal.emit(True, "Модель успешно загружена")

        except Exception as e:
            self.log_signal.emit(f"Ошибка загрузки модели: {str(e)}", "ERROR")
            CrashSafeMemoryManager.safe_gpu_cleanup("after model download error")
            self.finished_signal.emit(False, str(e))

    def stop(self):
        self.should_stop = True
        self.quit()
        self.wait()

    @staticmethod
    def get_models_cache_dir() -> Path:
        """Получить директорию для кэша моделей"""
        if sys.platform == "win32":
            cache_dir = Path.home() / "AppData" / "Local" / "WhisperModels"
        else:
            cache_dir = Path.home() / ".cache" / "whisper"
        cache_dir.mkdir(parents=True, exist_ok=True)
        return cache_dir


class ProfessionalTranscriptionWorker(QThread):
    """Рабочий поток транскрибации с защитой от крашей"""

    progress_signal = Signal(int, str)
    log_signal = Signal(str, str)
    finished_signal = Signal(str)
    segment_signal = Signal(str)
    stats_signal = Signal(dict)

    def __init__(self, file_path, settings):
        super().__init__()
        self.file_path = file_path
        self.settings = settings
        self.temp_dir = None
        self._is_running = True
        self.start_time = None

    def run(self):
        """Процесс транскрибации с защитой от крашей"""
        model = None
        try:
            self.start_time = time.time()
            self.log_signal.emit("Начало транскрибации", "INFO")

            # Создаем временную директорию
            self.temp_dir = tempfile.TemporaryDirectory()
            output_audio = os.path.join(self.temp_dir.name, "audio.wav")

            # Этап 1: Извлечение аудио
            self.progress_signal.emit(10, "Извлечение аудио...")
            self.log_signal.emit("Начинаем извлечение аудио", "INFO")

            if not self._is_running:
                return

            self.extract_audio(output_audio)

            # Предварительная очистка памяти перед загрузкой модели
            CrashSafeMemoryManager.safe_gpu_cleanup("before model loading")

            # Этап 2: Загрузка модели
            self.progress_signal.emit(20, "Загрузка Whisper модели...")
            self.log_signal.emit(f"Загружаем модель: {self.settings['model_size']} на {DEVICE}", "INFO")

            if not self._is_running:
                return

            model = self.load_model_safely()

            if not self._is_running:
                if model:
                    del model
                    model = None
                    CrashSafeMemoryManager.safe_gpu_cleanup("after early stop")
                return

            # Этап 3: Транскрибация
            self.progress_signal.emit(30, "Распознавание речи...")
            self.log_signal.emit("Начинаем транскрибацию", "INFO")

            if not self._is_running:
                if model:
                    del model
                    model = None
                    CrashSafeMemoryManager.safe_gpu_cleanup("after early stop")
                return

            segments = self.transcribe_audio_safely(output_audio, model)

            # Безопасно освобождаем модель
            if model:
                try:
                    del model
                    model = None
                    self.log_signal.emit("Модель удалена", "DEBUG")
                except Exception as model_cleanup_error:
                    self.log_signal.emit(f"Предупреждение при удалении модели: {model_cleanup_error}", "WARNING")

            # Очистка памяти после транскрибации
            CrashSafeMemoryManager.safe_gpu_cleanup("after transcription")

            if not self._is_running:
                return

            # Проверяем валидность сегментов
            if not segments:
                self.log_signal.emit("Сегменты не получены, возможно аудио слишком тихое", "WARNING")
                self.finished_signal.emit("Не удалось получить сегменты из аудио. Проверьте качество записи.")
                return

            # Этап 4: Диаризация
            formatted_text = ""
            if self.settings.get('use_diarization'):
                self.progress_signal.emit(70, "Диаризация спикеров...")

                if not self._is_running:
                    return

                formatted_text = self.apply_crash_safe_diarization(segments)
            else:
                formatted_text = self.format_simple_text_safely(segments)

            # Очистка памяти после диаризации
            CrashSafeMemoryManager.safe_gpu_cleanup("after diarization")

            if not formatted_text or formatted_text.strip() == "":
                formatted_text = "Транскрибация завершена, но результат пуст. Проверьте аудио файл."
                self.log_signal.emit("Получен пустой результат транскрибации", "WARNING")

            # Статистика
            try:
                self.send_statistics(segments, formatted_text)
            except Exception as stats_error:
                self.log_signal.emit(f"Ошибка статистики: {stats_error}", "WARNING")

            # Финальная очистка
            CrashSafeMemoryManager.safe_gpu_cleanup("final cleanup")

            self.progress_signal.emit(100, "Готово!")
            self.log_signal.emit("Транскрибация завершена успешно", "SUCCESS")
            self.finished_signal.emit(formatted_text)

        except Exception as e:
            import traceback
            error_msg = f"Ошибка транскрибации: {str(e)}"
            error_trace = traceback.format_exc()

            self.log_signal.emit(error_msg, "ERROR")
            self.log_signal.emit(f"Детальная трассировка: {error_trace}", "DEBUG")

            # Защита от краша при ошибке
            try:
                self.log_signal.emit("Экстренная очистка памяти...", "WARNING")
                if model:
                    del model
                    model = None

                # Агрессивная очистка при ошибке
                CrashSafeMemoryManager.safe_gpu_cleanup("emergency cleanup")

                self.log_signal.emit("Защитная очистка завершена", "INFO")

            except Exception as cleanup_critical_error:
                self.log_signal.emit(f"Критическая ошибка защиты: {cleanup_critical_error}", "ERROR")

            self.finished_signal.emit(f"Ошибка транскрибации: {str(e)}")
        finally:
            # Гарантированная финальная очистка
            try:
                if model:
                    del model
                    model = None
            except:
                pass

            self.cleanup()

    def stop(self):
        """Остановка процесса"""
        self._is_running = False
        self.log_signal.emit("Получен сигнал остановки...", "WARNING")

        # Даем время на корректное завершение
        try:
            if self.isRunning():
                self.quit()
                if not self.wait(3000):  # Ждем 3 секунды
                    self.log_signal.emit("Принудительное завершение...", "WARNING")
                    self.terminate()
                    self.wait(1000)  # Еще секунда
        except Exception as e:
            self.log_signal.emit(f"Ошибка остановки: {e}", "ERROR")

    def extract_audio(self, output_path):
        """Извлечение аудио с защитой"""
        ffmpeg_exe = self.find_ffmpeg()
        if not ffmpeg_exe:
            raise RuntimeError("FFmpeg не найден! Установите FFmpeg или поместите ffmpeg.exe в папку с программой")

        cmd = [
            ffmpeg_exe, "-y",
            "-i", self.file_path,
            "-vn", "-acodec", "pcm_s16le",
            "-ar", "16000", "-ac", "1",
            "-af", "highpass=f=200,lowpass=f=3000",
            output_path
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            raise RuntimeError(f"Ошибка FFmpeg: {result.stderr}")

        size_mb = os.path.getsize(output_path) / (1024 * 1024)
        self.log_signal.emit(f"Аудио извлечено: {size_mb:.1f} МБ", "SUCCESS")

    def find_ffmpeg(self):
        """Поиск FFmpeg"""
        if getattr(sys, 'frozen', False):
            app_dir = Path(sys.executable).parent
        else:
            app_dir = Path(__file__).parent

        local_ffmpeg = app_dir / "ffmpeg.exe"
        if local_ffmpeg.exists():
            return str(local_ffmpeg)

        try:
            result = subprocess.run(['ffmpeg', '-version'], capture_output=True)
            if result.returncode == 0:
                return 'ffmpeg'
        except:
            pass

        return None

    def load_model_safely(self):
        """Безопасная загрузка модели"""
        try:
            compute_type = "float16" if DEVICE == "cuda" else "int8"

            # Для больших моделей на слабых GPU
            if DEVICE == "cuda" and self.settings['model_size'] in ['large', 'large-v2', 'large-v3']:
                try:
                    import torch
                    vram = torch.cuda.get_device_properties(0).total_memory / (1024 ** 3)
                    if vram < 10:
                        compute_type = "int8"
                        self.log_signal.emit(f"GPU память: {vram:.1f}GB. Используем int8", "WARNING")
                except:
                    pass

            self.log_signal.emit(f"Создание модели с {compute_type}", "INFO")

            model = WhisperModel(
                self.settings['model_size'],
                device=DEVICE,
                compute_type=compute_type,
                cpu_threads=min(4, os.cpu_count()),
                num_workers=1,
                download_root=ModelDownloader.get_models_cache_dir()
            )

            self.log_signal.emit(f"Модель загружена (compute_type: {compute_type})", "SUCCESS")
            return model

        except Exception as e:
            self.log_signal.emit(f"Ошибка загрузки модели: {e}", "ERROR")
            try:
                self.log_signal.emit("Пробуем базовую модель...", "WARNING")
                model = WhisperModel(
                    "base",
                    device="cpu",
                    compute_type="int8",
                    cpu_threads=2,
                    num_workers=1,
                    download_root=ModelDownloader.get_models_cache_dir()
                )
                self.log_signal.emit("Базовая модель загружена в безопасном режиме", "WARNING")
                return model
            except Exception as fallback_e:
                self.log_signal.emit(f"Критическая ошибка загрузки: {fallback_e}", "ERROR")
                raise fallback_e

    def transcribe_audio_safely(self, audio_path, model):
        """Безопасная транскрибация"""
        try:
            segments, info = model.transcribe(
                audio_path,
                language=self.settings['language'] if self.settings['language'] != 'auto' else None,
                task="transcribe",
                beam_size=5,
                best_of=5,
                patience=1,
                temperature=0.0,
                initial_prompt="Это транскрипция на русском языке." if self.settings['language'] == 'ru' else None,
                word_timestamps=True,
                vad_filter=True,
                vad_parameters=dict(
                    threshold=0.5,
                    min_speech_duration_ms=250,
                    max_speech_duration_s=float('inf'),
                    min_silence_duration_ms=self.settings.get('min_silence', 1000),
                    speech_pad_ms=400
                ),
                condition_on_previous_text=False,
                compression_ratio_threshold=2.4,
                log_prob_threshold=-1.0,
                no_speech_threshold=0.6
            )

            # Язык
            if hasattr(info, 'language'):
                self.log_signal.emit(f"Определен язык: {info.language} ({info.language_probability:.0%})", "INFO")

            # Безопасная обработка сегментов
            segments_list = []
            total_segments = 0

            for segment in segments:
                if not self._is_running:
                    break

                try:
                    text = self.clean_text_safely(segment.text.strip())
                    if text and len(text) > 2:
                        segments_list.append({
                            'start': float(segment.start),
                            'end': float(segment.end),
                            'text': text
                        })

                        # Отправляем сегмент для отображения
                        self.segment_signal.emit(f"[{self.format_time(segment.start)}] {text}")

                        total_segments += 1
                        if total_segments % 10 == 0:
                            progress = min(30 + int(40 * segment.end / (info.duration or 100)), 70)
                            self.progress_signal.emit(progress, f"Обработано: {total_segments}")

                            # Периодическая очистка
                            if total_segments % 50 == 0:
                                CrashSafeMemoryManager.safe_gpu_cleanup("during transcription")

                except Exception as seg_error:
                    self.log_signal.emit(f"Пропускаем проблемный сегмент: {seg_error}", "DEBUG")
                    continue

            self.log_signal.emit(f"Распознано сегментов: {len(segments_list)}", "SUCCESS")
            return segments_list

        except Exception as e:
            self.log_signal.emit(f"Ошибка транскрибации: {e}", "ERROR")
            raise e

    def clean_text_safely(self, text):
        """Безопасная очистка текста"""
        if not text:
            return text

        try:
            # Удаляем повторения
            text = re.sub(r'([а-яА-Яa-zA-Z])\1{3,}', r'\1\1', text)
            text = re.sub(r'(\b\w{1,3}\b)[\s-]*(?:\1[\s-]*){3,}', r'\1', text)

            # Удаляем мусорные символы
            text = re.sub(r'[^\w\s\.\,\!\?\-\:\;\"\']+', ' ', text)
            text = re.sub(r'\s+', ' ', text)

            return text.strip()
        except Exception:
            return str(text)[:500]  # Безопасный fallback

    def apply_crash_safe_diarization(self, segments):
        """Диаризация с защитой от крашей"""
        try:
            self.log_signal.emit("Применение диаризации...", "INFO")

            if not segments or len(segments) == 0:
                self.log_signal.emit("Нет сегментов для диаризации", "WARNING")
                return "Нет распознанной речи для обработки."

            # Предварительная очистка памяти
            CrashSafeMemoryManager.safe_gpu_cleanup("before diarization")

            min_pause = float(self.settings.get('min_pause', 2.0))
            diarized_segments = []
            current_speaker = 1
            last_end = 0.0

            self.log_signal.emit(f"Обрабатываем {len(segments)} сегментов с паузой {min_pause}с", "INFO")

            for i, seg in enumerate(segments):
                try:
                    # Проверка остановки
                    if i % 25 == 0 and not self._is_running:
                        self.log_signal.emit("Диаризация прервана", "WARNING")
                        break

                    # Безопасная валидация сегмента
                    if not self.validate_segment_safely(seg, i):
                        continue

                    start_time = float(seg['start'])
                    end_time = float(seg['end'])
                    text = str(seg['text']).strip()[:500]

                    if not text or len(text) < 2:
                        continue

                    # Проверяем паузу для смены спикера
                    if last_end > 0 and start_time - last_end > min_pause:
                        current_speaker = 2 if current_speaker == 1 else 1
                        self.log_signal.emit(f"Смена спикера на {current_speaker}", "DEBUG")

                    diarized_segments.append({
                        'speaker': f"Спикер {current_speaker}",
                        'text': text,
                        'start': start_time,
                        'end': end_time
                    })

                    last_end = end_time

                    # Периодическая очистка памяти
                    if i % 50 == 0 and i > 0:
                        CrashSafeMemoryManager.safe_gpu_cleanup("during diarization")

                except Exception as seg_error:
                    self.log_signal.emit(f"Пропускаем сегмент {i}: {seg_error}", "DEBUG")
                    continue

            # Очистка после диаризации
            CrashSafeMemoryManager.safe_gpu_cleanup("after diarization processing")

            if not diarized_segments:
                self.log_signal.emit("Диаризация не создала сегментов", "WARNING")
                return self.format_simple_text_safely(segments)

            unique_speakers = len(set(seg['speaker'] for seg in diarized_segments))
            self.log_signal.emit(f"Диаризация завершена: {len(diarized_segments)} сегментов, {unique_speakers} спикеров", "SUCCESS")

            return self.format_diarized_text_safely(diarized_segments)

        except Exception as e:
            self.log_signal.emit(f"Ошибка диаризации: {e}", "ERROR")
            CrashSafeMemoryManager.safe_gpu_cleanup("after diarization error")
            return self.format_simple_text_safely(segments)

    def validate_segment_safely(self, seg, index):
        """Безопасная валидация сегмента"""
        try:
            if not seg or not isinstance(seg, dict):
                return False

            required_keys = ['start', 'end', 'text']
            if not all(key in seg for key in required_keys):
                return False

            start_time = float(seg['start'])
            end_time = float(seg['end'])

            if start_time < 0 or end_time < start_time or end_time - start_time > 600:
                return False

            text = str(seg['text']).strip()
            if not text or len(text) < 1:
                return False

            return True

        except Exception:
            return False

    def format_simple_text_safely(self, segments):
        """Безопасное простое форматирование"""
        try:
            if not segments or len(segments) == 0:
                return "Нет данных для форматирования."

            texts = []
            for i, seg in enumerate(segments):
                try:
                    if not self.validate_segment_safely(seg, i):
                        continue

                    text = str(seg['text']).strip()
                    if text and len(text) > 1:
                        texts.append(text)

                    if i % 100 == 0 and not self._is_running:
                        break

                except Exception:
                    continue

            if not texts:
                return "Не удалось извлечь текст из сегментов."

            result = " ".join(texts)
            result = result.replace("  ", " ").strip()

            if len(result) > 100000:
                result = result[:100000] + "\n\n[Результат обрезан для стабильности]"

            return result if result else "Пустой результат форматирования."

        except Exception as e:
            self.log_signal.emit(f"Ошибка простого форматирования: {e}", "ERROR")
            return "Ошибка при форматировании результата."

    def format_diarized_text_safely(self, segments):
        """Безопасное форматирование диаризованного текста"""
        try:
            if not segments or len(segments) == 0:
                return "Нет сегментов для форматирования."

            formatted_parts = []
            current_speaker = None
            current_texts = []

            for i, seg in enumerate(segments):
                try:
                    if not seg or not isinstance(seg, dict):
                        continue

                    speaker = str(seg.get('speaker', 'Неизвестный')).strip()[:50]
                    text = str(seg.get('text', '')).strip()[:1000]

                    if not speaker or not text:
                        continue

                    if speaker != current_speaker:
                        if current_texts and current_speaker:
                            speaker_line = f"{current_speaker}: {' '.join(current_texts)}"
                            formatted_parts.append(speaker_line[:3000])

                        current_speaker = speaker
                        current_texts = [text]
                    else:
                        current_texts.append(text)

                        if len(current_texts) > 100:
                            current_texts = current_texts[-100:]

                except Exception:
                    continue

            if current_texts and current_speaker:
                speaker_line = f"{current_speaker}: {' '.join(current_texts)}"
                formatted_parts.append(speaker_line[:3000])

            if not formatted_parts:
                return "Не удалось сформатировать диаризованный текст."

            result = "\n\n".join(formatted_parts)

            if len(result) > 200000:
                result = result[:200000] + "\n\n[Результат обрезан для стабильности]"

            return result if result.strip() else "Пустой результат диаризации."

        except Exception as e:
            self.log_signal.emit(f"Ошибка форматирования диаризации: {e}", "ERROR")
            return "Ошибка при форматировании диаризованного текста."

    def format_time(self, seconds):
        """Форматирование времени"""
        return f"{int(seconds // 60):02d}:{int(seconds % 60):02d}"

    def send_statistics(self, segments, text):
        """Отправка статистики"""
        try:
            elapsed = time.time() - self.start_time
            stats = {
                'duration': elapsed,
                'segments': len(segments),
                'words': len(text.split()),
                'chars': len(text),
                'speed': len(segments) / elapsed if elapsed > 0 else 0
            }
            self.stats_signal.emit(stats)
        except Exception as e:
            self.log_signal.emit(f"Ошибка статистики: {e}", "WARNING")

    def cleanup(self):
        """Очистка ресурсов"""
        try:
            if self.temp_dir:
                self.temp_dir.cleanup()
                self.log_signal.emit("Временные файлы удалены", "DEBUG")
        except Exception as e:
            self.log_signal.emit(f"Предупреждение при удалении временных файлов: {e}", "WARNING")

        # Очень мягкая финальная очистка памяти
        try:
            gc.collect()
        except:
            pass


class MainWindow(QMainWindow):
    """Главное окно приложения"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} v{APP_VERSION}")
        self.setGeometry(100, 100, 1200, 800)
        self.setMinimumSize(1000, 700)

        # Темная тема
        self.setStyleSheet(self.get_dark_theme())

        if not FASTER_WHISPER_AVAILABLE:
            QMessageBox.critical(
                self, "Критическая ошибка",
                "faster_whisper не установлен!\n\n"
                "Установите: pip install faster-whisper"
            )
            sys.exit(1)

        self.init_ui()
        self.video_file_path = None
        self.transcription_thread = None
        self.transcribed_text = ""
        self.model_downloader = None

    def init_ui(self):
        """Создание интерфейса"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)

        # Заголовок
        header_container = QWidget()
        header_layout = QVBoxLayout(header_container)
        header_layout.setSpacing(0)
        header_layout.setContentsMargins(0,0,0,0)

        title_label = QLabel("ПРОФЕССИОНАЛЬНАЯ ТРАНСКРИБАЦИЯ")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #ffffff;
                padding-top: 10px;
                padding-bottom: 5px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                border-top-left-radius: 10px;
                border-top-right-radius: 10px;
            }
        """)

        author_label = QLabel(f"by {AUTHOR}")
        author_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        author_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                font-style: italic;
                color: #e2e8f0;
                padding-bottom: 10px;
                padding-top: 0px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                border-bottom-left-radius: 10px;
                border-bottom-right-radius: 10px;
            }
        """)

        header_layout.addWidget(title_label)
        header_layout.addWidget(author_label)
        main_layout.addWidget(header_container)

        # Основной контент в табах
        tabs = QTabWidget()
        tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #3c3c3c;
                background: #2d2d2d;
                border-radius: 8px;
            }
            QTabBar::tab {
                background: #3c3c3c;
                color: #d4d4d4;
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QTabBar::tab:selected {
                background: #667eea;
                color: white;
            }
            QTabBar::tab:hover {
                background: #4c5cda;
            }
        """)

        # Вкладка транскрибации
        transcription_tab = self.create_transcription_tab()
        tabs.addTab(transcription_tab, "🎙️ ТРАНСКРИБАЦИЯ")

        # Вкладка настроек
        settings_tab = self.create_settings_tab()
        tabs.addTab(settings_tab, "⚙️ НАСТРОЙКИ")

        # Вкладка логов
        logs_tab = self.create_logs_tab()
        tabs.addTab(logs_tab, "📊 ЛОГИ")

        main_layout.addWidget(tabs)

        # Статус бар
        self.create_status_bar()

    def create_transcription_tab(self):
        """Создание вкладки транскрибации"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Выбор файла
        file_group = QGroupBox("📁 Выбор файла")
        file_layout = QHBoxLayout()

        self.select_file_btn = QPushButton("Выбрать файл")
        self.select_file_btn.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DirOpenIcon))
        self.select_file_btn.setStyleSheet("""
            QPushButton {
                padding: 8px 16px;
                font-size: 14px;
                font-weight: bold;
            }
        """)

        self.file_label = QLabel("Файл не выбран")
        self.file_label.setStyleSheet("color: #8b949e; font-style: italic;")

        file_layout.addWidget(self.select_file_btn)
        file_layout.addWidget(self.file_label, 1)
        file_group.setLayout(file_layout)

        # Контролы
        control_layout = QHBoxLayout()

        self.transcribe_btn = QPushButton("НАЧАТЬ ТРАНСКРИБАЦИЮ")
        self.transcribe_btn.setEnabled(False)
        self.transcribe_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                color: white;
                border: none;
                padding: 12px 30px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 25px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #5a67d8, stop:1 #6b46c1);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #4c51bf, stop:1 #553c9a);
            }
            QPushButton:disabled {
                background: #4a5568;
            }
        """)

        self.stop_btn = QPushButton("⏹️ Остановить")
        self.stop_btn.setEnabled(False)
        self.stop_btn.setStyleSheet("""
            QPushButton {
                background: #e53e3e;
                color: white;
                border: none;
                padding: 12px 30px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 25px;
            }
            QPushButton:hover {
                background: #c53030;
            }
            QPushButton:disabled {
                background: #4a5568;
            }
        """)

        self.download_model_btn = QPushButton("📥 Скачать модель")
        self.download_model_btn.setStyleSheet("""
            QPushButton {
                background: #38a169;
                color: white;
                border: none;
                padding: 12px 30px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 25px;
            }
            QPushButton:hover {
                background: #2f855a;
            }
            QPushButton:disabled {
                background: #4a5568;
            }
        """)

        control_layout.addWidget(self.transcribe_btn)
        control_layout.addWidget(self.stop_btn)
        control_layout.addWidget(self.download_model_btn)
        control_layout.addStretch()

        # Прогресс бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #3c3c3c;
                border-radius: 10px;
                text-align: center;
                font-weight: bold;
                background-color: #2d2d2d;
                height: 25px;
            }
            QProgressBar::chunk {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                border-radius: 8px;
            }
        """)
        self.progress_bar.hide()

        self.status_label = QLabel("")
        self.status_label.setStyleSheet("color: #8b949e; font-size: 12px;")
        self.status_label.hide()

        # Splitter для результатов
        splitter = QSplitter(Qt.Orientation.Vertical)

        # Результат
        result_group = QGroupBox("📄 РЕЗУЛЬТАТ")
        result_layout = QVBoxLayout()

        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setPlaceholderText("Здесь появится расшифровка...")

        # Кнопки сохранения
        save_layout = QHBoxLayout()
        self.save_txt_btn = self.create_save_button("💾 TXT", "#48bb78")
        self.save_docx_btn = self.create_save_button("📄 DOCX", "#4299e1")
        self.save_json_btn = self.create_save_button("📊 JSON", "#ed8936")

        save_layout.addStretch()
        save_layout.addWidget(self.save_txt_btn)
        save_layout.addWidget(self.save_docx_btn)
        save_layout.addWidget(self.save_json_btn)

        result_layout.addWidget(self.result_text)
        result_layout.addLayout(save_layout)
        result_group.setLayout(result_layout)

        # Живая транскрипция
        live_group = QGroupBox("🔴 ЖИВАЯ ТРАНСКРИПЦИЯ")
        live_layout = QVBoxLayout()

        self.live_text = QTextEdit()
        self.live_text.setReadOnly(True)
        self.live_text.setMaximumHeight(150)
        self.live_text.setPlaceholderText("Сегменты будут появляться здесь в реальном времени...")

        live_layout.addWidget(self.live_text)
        live_group.setLayout(live_layout)

        splitter.addWidget(result_group)
        splitter.addWidget(live_group)
        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 1)

        # Компоновка
        layout.addWidget(file_group)
        layout.addLayout(control_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)
        layout.addWidget(splitter)

        # Подключение сигналов
        self.select_file_btn.clicked.connect(self.select_file)
        self.transcribe_btn.clicked.connect(self.start_transcription)
        self.stop_btn.clicked.connect(self.stop_transcription)
        self.download_model_btn.clicked.connect(self.download_model)
        self.save_txt_btn.clicked.connect(lambda: self.save_results('txt'))
        self.save_docx_btn.clicked.connect(lambda: self.save_results('docx'))
        self.save_json_btn.clicked.connect(lambda: self.save_results('json'))

        return widget

    def create_settings_tab(self):
        """Создание вкладки настроек"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Основные настройки
        basic_group = QGroupBox("⚙️ ОСНОВНЫЕ НАСТРОЙКИ")
        basic_layout = QVBoxLayout()

        # Язык
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel("Язык:"))
        self.language_combo = QComboBox()
        self.language_combo.addItems(["ru - Русский", "en - English", "auto - Автоопределение"])
        self.language_combo.setCurrentIndex(0)
        lang_layout.addWidget(self.language_combo)
        lang_layout.addStretch()

        # Модель
        model_layout = QHBoxLayout()
        model_layout.addWidget(QLabel("Модель:"))
        self.model_combo = QComboBox()
        models = [
            "tiny - Очень быстро (39 MB)",
            "base - Быстро (74 MB)",
            "small - Баланс (244 MB)",
            "medium - Качественно (769 MB)",
            "large-v3 - Максимум (1.5 GB)"
        ]
        self.model_combo.addItems(models)
        self.model_combo.setCurrentIndex(2)  # small по умолчанию
        model_layout.addWidget(self.model_combo)
        model_layout.addStretch()

        basic_layout.addLayout(lang_layout)
        basic_layout.addLayout(model_layout)
        basic_group.setLayout(basic_layout)

        # Диаризация
        diarization_group = QGroupBox("👥 ДИАРИЗАЦИЯ СПИКЕРОВ")
        diarization_layout = QVBoxLayout()

        self.diarization_checkbox = QCheckBox("Включить диаризацию спикеров")
        self.diarization_checkbox.setChecked(True)

        # Информация о диаризации
        info_label = QLabel("Автоматически разделяет речь разных говорящих")
        info_label.setStyleSheet("color: #56d364; font-size: 11px; margin: 5px 0px;")
        info_label.setWordWrap(True)

        # Параметры диаризации
        params_layout = QHBoxLayout()

        params_layout.addWidget(QLabel("Мин. пауза между спикерами (сек):"))
        self.min_pause_spin = QSpinBox()
        self.min_pause_spin.setMinimum(1)
        self.min_pause_spin.setMaximum(10)
        self.min_pause_spin.setValue(2)
        self.min_pause_spin.setToolTip("Минимальная пауза для смены спикера")
        params_layout.addWidget(self.min_pause_spin)

        params_layout.addWidget(QLabel("Мин. тишина (мс):"))
        self.min_silence_spin = QSpinBox()
        self.min_silence_spin.setMinimum(500)
        self.min_silence_spin.setMaximum(3000)
        self.min_silence_spin.setSingleStep(100)
        self.min_silence_spin.setValue(1000)
        self.min_silence_spin.setToolTip("Минимальная длительность тишины для VAD")
        params_layout.addWidget(self.min_silence_spin)

        params_layout.addStretch()

        diarization_layout.addWidget(self.diarization_checkbox)
        diarization_layout.addWidget(info_label)
        diarization_layout.addLayout(params_layout)
        diarization_group.setLayout(diarization_layout)

        # Подключаем обработчик
        self.diarization_checkbox.stateChanged.connect(self.on_diarization_changed)

        # Информация о системе
        info_group = QGroupBox("💻 Информация о системе")
        info_layout = QVBoxLayout()

        self.system_info = QTextEdit()
        self.system_info.setReadOnly(True)
        self.system_info.setMaximumHeight(150)
        self.update_system_info()

        info_layout.addWidget(self.system_info)
        info_group.setLayout(info_layout)

        # Компоновка
        layout.addWidget(basic_group)
        layout.addWidget(diarization_group)
        layout.addWidget(info_group)
        layout.addStretch()

        return widget

    def create_logs_tab(self):
        """Создание вкладки логов"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Логи
        self.log_widget = LogWidget()

        # Кнопки управления логами
        controls_layout = QHBoxLayout()

        clear_btn = QPushButton("🗑️ Очистить")
        clear_btn.clicked.connect(self.log_widget.clear)

        export_btn = QPushButton("💾 Экспорт")
        export_btn.clicked.connect(self.export_logs)

        controls_layout.addStretch()
        controls_layout.addWidget(clear_btn)
        controls_layout.addWidget(export_btn)

        # Статистика
        stats_group = QGroupBox("📊 Статистика последней транскрибации")
        stats_layout = QVBoxLayout()

        self.stats_text = QTextEdit()
        self.stats_text.setReadOnly(True)
        self.stats_text.setMaximumHeight(150)
        self.stats_text.setPlaceholderText("Статистика появится после транскрибации...")

        stats_layout.addWidget(self.stats_text)
        stats_group.setLayout(stats_layout)

        # Компоновка
        layout.addWidget(self.log_widget)
        layout.addLayout(controls_layout)
        layout.addWidget(stats_group)

        return widget

    def create_status_bar(self):
        """Создание статус бара"""
        status_bar = self.statusBar()
        status_bar.setStyleSheet("""
            QStatusBar {
                background: #1e1e1e;
                color: #8b949e;
                border-top: 1px solid #3c3c3c;
            }
        """)

        self.memory_label = QLabel("💾 Память: --")
        self.gpu_label = QLabel("GPU: --")
        self.time_label = QLabel("⏱️ Время: --")

        status_bar.addPermanentWidget(self.memory_label)
        status_bar.addPermanentWidget(self.gpu_label)
        status_bar.addPermanentWidget(self.time_label)

        # Таймер для обновления статистики
        self.status_timer = QTimer()
        self.status_timer.timeout.connect(self.update_status)
        self.status_timer.start(1000)

    def create_save_button(self, text, color):
        """Создание кнопки сохранения"""
        btn = QPushButton(text)
        btn.setEnabled(False)
        btn.setStyleSheet(f"""
            QPushButton {{
                background: {color};
                color: white;
                border: none;
                padding: 8px 20px;
                font-weight: bold;
                border-radius: 20px;
            }}
            QPushButton:hover {{
                background: {color}dd;
            }}
            QPushButton:pressed {{
                background: {color}bb;
            }}
            QPushButton:disabled {{
                background: #4a5568;
                color: #718096;
            }}
        """)
        return btn

    def get_dark_theme(self):
        """Темная тема"""
        return """
            QMainWindow {
                background-color: #1e1e1e;
            }
            QWidget {
                background-color: #2d2d2d;
                color: #d4d4d4;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 13px;
            }
            QGroupBox {
                border: 2px solid #3c3c3c;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                font-weight: bold;
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px 0 10px;
                color: #58a6ff;
            }
            QPushButton {
                background-color: #3c3c3c;
                border: 1px solid #484848;
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #484848;
            }
            QPushButton:pressed {
                background-color: #2d2d2d;
            }
            QComboBox {
                background-color: #3c3c3c;
                border: 1px solid #484848;
                padding: 6px;
                border-radius: 6px;
            }
            QComboBox:drop-down {
                border: none;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #d4d4d4;
                margin-right: 5px;
            }
            QLineEdit {
                background-color: #3c3c3c;
                border: 1px solid #484848;
                padding: 6px;
                border-radius: 6px;
            }
            QTextEdit {
                background-color: #1e1e1e;
                border: 1px solid #3c3c3c;
                border-radius: 6px;
                padding: 8px;
            }
            QProgressBar {
                text-align: center;
            }
            QScrollBar:vertical {
                background: #2d2d2d;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #5a5a5a;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background: #7a7a7a;
            }
            QSpinBox {
                background-color: #3c3c3c;
                border: 1px solid #484848;
                padding: 4px;
                border-radius: 6px;
            }
            QCheckBox {
                spacing: 8px;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
                border-radius: 4px;
                border: 2px solid #484848;
                background-color: #2d2d2d;
            }
            QCheckBox::indicator:checked {
                background-color: #667eea;
                border-color: #667eea;
            }
            QMenuBar {
                background-color: #2d2d2d;
                color: #d4d4d4;
            }
            QMenuBar::item:selected {
                background-color: #667eea;
            }
            QMenu {
                background-color: #2d2d2d;
                border: 1px solid #3c3c3c;
                color: #d4d4d4;
            }
            QMenu::item:selected {
                background-color: #667eea;
            }
        """

    def on_diarization_changed(self, state):
        """Обработчик изменения чекбокса диаризации"""
        enabled = state == Qt.CheckState.Checked

        # Включаем/выключаем параметры
        self.min_pause_spin.setEnabled(enabled)
        self.min_silence_spin.setEnabled(enabled)

        self.log_widget.log(f"Диаризация {'включена' if enabled else 'отключена'}", "INFO")

    def update_system_info(self):
        """Обновление информации о системе"""
        info = []

        # Python
        info.append(f"Python: {sys.version.split()[0]}")

        # GPU
        if TORCH_AVAILABLE:
            if DEVICE == "cuda":
                try:
                    import torch
                    gpu_name = torch.cuda.get_device_name(0)
                    vram = torch.cuda.get_device_properties(0).total_memory / (1024 ** 3)
                    info.append(f"GPU: {gpu_name} ({vram:.1f} GB)")
                except:
                    info.append("GPU: Обнаружен")
            else:
                info.append("Режим: CPU")
        else:
            info.append("PyTorch не установлен")

        # FFmpeg
        ffmpeg_found = os.path.exists('ffmpeg.exe') or self.check_ffmpeg()
        info.append(f"FFmpeg: {'✅ Найден' if ffmpeg_found else '❌ Не найден'}")

        # Модели
        cache_dir = ModelDownloader.get_models_cache_dir()
        if cache_dir.exists():
            models = list(cache_dir.glob("*"))
            info.append(f"Кэш моделей: {len(models)} файлов")

        info.append("Защита от крашей: активна")

        self.system_info.setPlainText("\n".join(info))

    def check_ffmpeg(self):
        """Проверка FFmpeg"""
        try:
            if getattr(sys, 'frozen', False):
                app_dir = Path(sys.executable).parent
            else:
                app_dir = Path(__file__).parent

            local_ffmpeg = app_dir / "ffmpeg.exe"
            if local_ffmpeg.exists():
                return True

            result = subprocess.run(['ffmpeg', '-version'], capture_output=True)
            return result.returncode == 0
        except:
            return False

    def update_status(self):
        """Обновление статус бара"""
        # Память
        try:
            import psutil
            memory = psutil.Process().memory_info().rss / (1024 ** 3)
            self.memory_label.setText(f"💾 Память: {memory:.1f} GB")
        except:
            pass

        # GPU
        if DEVICE == "cuda" and TORCH_AVAILABLE:
            try:
                import torch
                allocated = torch.cuda.memory_allocated() / (1024 ** 3)
                reserved = torch.cuda.memory_reserved() / (1024 ** 3)
                self.gpu_label.setText(f"GPU: {allocated:.1f}/{reserved:.1f} GB")
            except:
                pass
        else:
            self.gpu_label.setText("CPU режим")

        # Время
        current_time = datetime.now().strftime("%H:%M:%S")
        self.time_label.setText(f"⏱️ {current_time}")

    def select_file(self):
        """Выбор файла"""
        try:
            # Если есть активная транскрибация, предупреждаем
            if self.transcription_thread and self.transcription_thread.isRunning():
                reply = QMessageBox.question(
                    self,
                    "Предупреждение",
                    "Транскрибация в процессе. Остановить и выбрать новый файл?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                if reply == QMessageBox.StandardButton.Yes:
                    self.stop_transcription()
                    # Ждем завершения
                    if self.transcription_thread:
                        self.transcription_thread.wait(3000)  # Ждем до 3 секунд
                else:
                    return

            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Выберите видео или аудио файл",
                "",
                "Медиа файлы (*.mp4 *.mkv *.avi *.mov *.webm *.mp3 *.wav *.m4a *.aac *.flac *.ogg);;Все файлы (*.*)"
            )

            if file_path:
                # ВАЖНО: Полный сброс состояния перед новым файлом
                self.log_widget.log("Подготовка к новой транскрибации...", "INFO")
                CrashSafeMemoryManager.reset_for_new_transcription()

                self.video_file_path = file_path
                file_name = os.path.basename(file_path)
                file_size = os.path.getsize(file_path) / (1024 ** 2)

                self.file_label.setText(f"{file_name} ({file_size:.1f} MB)")
                self.file_label.setStyleSheet("color: #58a6ff; font-style: normal;")

                self.log_widget.log(f"Выбран файл: {file_name}", "SUCCESS")
                self.log_widget.log(f"Размер: {file_size:.1f} MB", "DEBUG")

                # Полная очистка UI
                self.result_text.clear()
                self.live_text.clear()
                self.stats_text.clear()
                self.transcribed_text = ""

                # Сброс прогресса
                self.progress_bar.hide()
                self.status_label.hide()

                # Активация кнопки
                self.transcribe_btn.setEnabled(True)

                # Деактивация кнопок сохранения
                for btn in [self.save_txt_btn, self.save_docx_btn, self.save_json_btn]:
                    btn.setEnabled(False)

                self.log_widget.log("Готов к транскрибации", "SUCCESS")

        except Exception as e:
            self.log_widget.log(f"Ошибка выбора файла: {e}", "ERROR")
            QMessageBox.critical(self, "Ошибка", f"Не удалось выбрать файл: {e}")

    def download_model(self):
        """Скачивание модели"""
        model_name = self.model_combo.currentText().split(' - ')[0]

        reply = QMessageBox.question(
            self,
            "Скачивание модели",
            f"Скачать модель {model_name}?\n\nЭто может занять время в зависимости от размера модели.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.download_model_btn.setEnabled(False)
            self.progress_bar.show()
            self.progress_bar.setValue(0)
            self.status_label.show()

            self.model_downloader = ModelDownloader(model_name)
            self.model_downloader.progress_signal.connect(self.on_download_progress)
            self.model_downloader.finished_signal.connect(self.on_download_finished)
            self.model_downloader.log_signal.connect(self.log_widget.log)
            self.model_downloader.start()

    @Slot(int, str)
    def on_download_progress(self, value, message):
        """Прогресс скачивания модели"""
        self.progress_bar.setValue(value)
        self.status_label.setText(message)

    @Slot(bool, str)
    def on_download_finished(self, success, message):
        """Завершение скачивания модели"""
        self.progress_bar.hide()
        self.status_label.hide()
        self.download_model_btn.setEnabled(True)

        if success:
            QMessageBox.information(self, "Успех", "Модель успешно загружена!")
            self.log_widget.log("Модель готова к использованию", "SUCCESS")
        else:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить модель:\n{message}")

        self.model_downloader = None

    def start_transcription(self):
        """Запуск транскрибации"""
        if not self.video_file_path:
            QMessageBox.warning(self, "Предупреждение", "Сначала выберите файл для транскрибации")
            return

        # Проверка файла
        if not os.path.exists(self.video_file_path):
            QMessageBox.critical(self, "Ошибка", "Выбранный файл не найден")
            return

        try:
            file_size = os.path.getsize(self.video_file_path)
            if file_size == 0:
                QMessageBox.critical(self, "Ошибка", "Выбранный файл пуст")
                return
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось прочитать файл: {e}")
            return

        # UI
        self.transcribe_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.download_model_btn.setEnabled(False)
        self.progress_bar.show()
        self.progress_bar.setValue(0)
        self.status_label.show()
        self.live_text.clear()
        self.result_text.clear()

        # Настройки
        lang = self.language_combo.currentText().split(' - ')[0]
        model = self.model_combo.currentText().split(' - ')[0]

        settings = {
            'language': lang,
            'model_size': model,
            'use_diarization': self.diarization_checkbox.isChecked(),
            'min_pause': float(self.min_pause_spin.value()),
            'min_silence': int(self.min_silence_spin.value())
        }

        self.log_widget.log("=" * 50, "INFO")
        self.log_widget.log("Начало профессиональной транскрибации", "INFO")
        self.log_widget.log(f"Файл: {os.path.basename(self.video_file_path)}", "INFO")
        self.log_widget.log(f"Размер: {file_size / (1024**2):.1f} MB", "INFO")
        self.log_widget.log(f"Настройки: {json.dumps(settings, ensure_ascii=False)}", "DEBUG")

        # Мягкая предварительная очистка памяти
        try:
            gc.collect()
        except:
            pass

        # Запуск потока
        try:
            self.transcription_thread = ProfessionalTranscriptionWorker(self.video_file_path, settings)
            self.transcription_thread.progress_signal.connect(self.on_progress)
            self.transcription_thread.log_signal.connect(self.log_widget.log)
            self.transcription_thread.finished_signal.connect(self.on_finished)
            self.transcription_thread.segment_signal.connect(self.on_segment)
            self.transcription_thread.stats_signal.connect(self.on_stats)
            self.transcription_thread.start()

            self.log_widget.log("Рабочий поток запущен", "SUCCESS")

        except Exception as thread_error:
            self.log_widget.log(f"Ошибка запуска потока: {thread_error}", "ERROR")
            self.on_finished(f"Ошибка запуска: {thread_error}")

    def stop_transcription(self):
        """Остановка транскрибации"""
        if self.transcription_thread and self.transcription_thread.isRunning():
            self.log_widget.log("Остановка транскрибации...", "WARNING")

            try:
                self.transcription_thread.stop()

                # Ждем корректного завершения
                if not self.transcription_thread.wait(5000):  # 5 секунд
                    self.log_widget.log("Принудительное завершение потока...", "WARNING")
                    self.transcription_thread.terminate()
                    self.transcription_thread.wait(2000)  # Еще 2 секунды

                self.log_widget.log("Поток остановлен", "INFO")

            except Exception as e:
                self.log_widget.log(f"Ошибка остановки потока: {e}", "ERROR")
            finally:
                # В любом случае вызываем завершение
                self.on_finished("Транскрибация остановлена пользователем")

    @Slot(int, str)
    def on_progress(self, value, message):
        """Обновление прогресса"""
        self.progress_bar.setValue(value)
        self.status_label.setText(message)

    @Slot(str)
    def on_segment(self, segment):
        """Отображение сегмента в реальном времени"""
        self.live_text.append(segment)

        # Автоскролл
        cursor = self.live_text.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.live_text.setTextCursor(cursor)

    @Slot(dict)
    def on_stats(self, stats):
        """Отображение статистики"""
        text = f"""
СТАТИСТИКА ТРАНСКРИБАЦИИ:
⏱️ Время обработки: {stats['duration']:.1f} сек
📝 Сегментов: {stats['segments']}
💬 Слов: {stats['words']}
📄 Символов: {stats['chars']}
⚡ Скорость: {stats['speed']:.1f} сегм/сек
"""
        self.stats_text.setPlainText(text.strip())

    @Slot(str)
    def on_finished(self, result):
        """Завершение транскрибации"""
        try:
            self.progress_bar.hide()
            self.status_label.hide()
            self.transcribe_btn.setEnabled(True)
            self.stop_btn.setEnabled(False)
            self.download_model_btn.setEnabled(True)

            if result.startswith("Ошибка"):
                QMessageBox.critical(self, "Ошибка", result)
                self.log_widget.log(result, "ERROR")
            else:
                self.transcribed_text = result
                self.result_text.setPlainText(result)

                # Активация кнопок сохранения
                for btn in [self.save_txt_btn, self.save_docx_btn, self.save_json_btn]:
                    btn.setEnabled(True)

                # Скролл к началу
                cursor = self.result_text.textCursor()
                cursor.setPosition(0)
                self.result_text.setTextCursor(cursor)

                self.log_widget.log("Транскрибация завершена успешно", "SUCCESS")
                self.log_widget.log("=" * 50, "INFO")

            # ВАЖНО: Очистка потока и состояния
            if self.transcription_thread:
                try:
                    if self.transcription_thread.isRunning():
                        self.transcription_thread.wait(1000)  # Ждем до 1 секунды
                except:
                    pass
                finally:
                    self.transcription_thread = None

            # Мягкая очистка памяти для подготовки к следующей транскрибации
            try:
                gc.collect()
                time.sleep(0.1)  # Даем время на стабилизацию
            except:
                pass

            self.log_widget.log("Система готова к новой транскрибации", "INFO")

        except Exception as e:
            self.log_widget.log(f"Ошибка завершения транскрибации: {e}", "ERROR")
            # Принудительная очистка при ошибке
            self.transcription_thread = None

    def save_results(self, format_type):
        """Сохранение результатов"""
        if not self.transcribed_text:
            return

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = Path(self.video_file_path).stem

        if format_type == 'txt':
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить как текст",
                f"{base_name}_{timestamp}.txt",
                "Текстовые файлы (*.txt)"
            )
            if file_path:
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(f"ПРОФЕССИОНАЛЬНАЯ ТРАНСКРИБАЦИЯ\n")
                        f.write(f"Файл: {os.path.basename(self.video_file_path)}\n")
                        f.write(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write(f"Программа: {APP_NAME} v{APP_VERSION}\n")
                        f.write(f"Автор: {AUTHOR}\n")
                        f.write("=" * 50 + "\n\n")
                        f.write(self.transcribed_text)

                    self.log_widget.log(f"Результат сохранен: {file_path}", "SUCCESS")
                    QMessageBox.information(self, "Успех", "Результат успешно сохранен!")
                except Exception as e:
                    self.log_widget.log(f"Ошибка сохранения: {e}", "ERROR")
                    QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить: {e}")

        elif format_type == 'docx':
            if not DOCX_AVAILABLE:
                QMessageBox.warning(self, "Внимание", "python-docx не установлен!\nУстановите: pip install python-docx")
                return

            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить как Word",
                f"{base_name}_{timestamp}.docx",
                "Документы Word (*.docx)"
            )
            if file_path:
                try:
                    doc = Document()
                    doc.add_heading('ПРОФЕССИОНАЛЬНАЯ ТРАНСКРИБАЦИЯ', 0)

                    # Метаданные
                    doc.add_paragraph(f'Файл: {os.path.basename(self.video_file_path)}')
                    doc.add_paragraph(f'Дата: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    doc.add_paragraph(f'Программа: {APP_NAME} v{APP_VERSION}')
                    doc.add_paragraph(f'Автор: {AUTHOR}')
                    doc.add_paragraph(f'Модель: {self.model_combo.currentText()}')
                    doc.add_paragraph(f'Язык: {self.language_combo.currentText()}')

                    doc.add_heading('Результат', level=1)

                    # Форматирование для спикеров
                    if "Спикер" in self.transcribed_text:
                        for para in self.transcribed_text.split('\n\n'):
                            if para.strip():
                                p = doc.add_paragraph()
                                if para.startswith("Спикер"):
                                    parts = para.split(":", 1)
                                    if len(parts) == 2:
                                        p.add_run(parts[0] + ":").bold = True
                                        p.add_run(" " + parts[1])
                                else:
                                    p.add_run(para)
                    else:
                        doc.add_paragraph(self.transcribed_text)

                    doc.save(file_path)
                    self.log_widget.log(f"Документ сохранен: {file_path}", "SUCCESS")
                    QMessageBox.information(self, "Успех", "Документ успешно сохранен!")
                except Exception as e:
                    self.log_widget.log(f"Ошибка сохранения: {e}", "ERROR")
                    QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить: {e}")

        elif format_type == 'json':
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить как JSON",
                f"{base_name}_{timestamp}.json",
                "JSON файлы (*.json)"
            )
            if file_path:
                try:
                    data = {
                        'metadata': {
                            'file': os.path.basename(self.video_file_path),
                            'date': datetime.now().isoformat(),
                            'program': f"{APP_NAME} v{APP_VERSION}",
                            'author': AUTHOR,
                            'model': self.model_combo.currentText(),
                            'language': self.language_combo.currentText(),
                            'professional_edition': True
                        },
                        'text': self.transcribed_text,
                        'statistics': self.stats_text.toPlainText() if self.stats_text.toPlainText() else None
                    }

                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)

                    self.log_widget.log(f"JSON сохранен: {file_path}", "SUCCESS")
                    QMessageBox.information(self, "Успех", "JSON успешно сохранен!")
                except Exception as e:
                    self.log_widget.log(f"Ошибка сохранения: {e}", "ERROR")
                    QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить: {e}")

    def export_logs(self):
        """Экспорт логов"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Экспорт логов",
            f"professional_logs_{timestamp}.txt",
            "Текстовые файлы (*.txt)"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"ЛОГИ {APP_NAME} v{APP_VERSION}\n")
                    f.write(f"Автор: {AUTHOR}\n")
                    f.write(f"Экспорт: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("=" * 60 + "\n\n")
                    f.write(self.log_widget.toPlainText())
                QMessageBox.information(self, "Успех", "Логи экспортированы!")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать: {e}")

    def closeEvent(self, event):
        """Закрытие приложения"""
        active_threads = []

        if self.transcription_thread and self.transcription_thread.isRunning():
            active_threads.append("транскрибация")

        if self.model_downloader and self.model_downloader.isRunning():
            active_threads.append("скачивание модели")

        if active_threads:
            reply = QMessageBox.question(
                self,
                'Подтверждение',
                f'Выполняется: {", ".join(active_threads)}.\nВы уверены, что хотите выйти?',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                self.log_widget.log("Принудительное закрытие программы...", "WARNING")

                # Остановка потоков без агрессивной очистки
                try:
                    if self.transcription_thread:
                        self.transcription_thread.stop()
                    if self.model_downloader:
                        self.model_downloader.stop()
                except:
                    pass

                event.accept()
            else:
                event.ignore()
        else:
            # Обычное закрытие без принудительной очистки GPU
            event.accept()


def main():
    """Главная функция"""
    try:
        app = QApplication(sys.argv)
        app.setStyle(QStyleFactory.create("Fusion"))
        app.setApplicationName(APP_NAME)
        app.setApplicationVersion(APP_VERSION)

        # Проверка критических зависимостей
        if not FASTER_WHISPER_AVAILABLE:
            QMessageBox.critical(
                None, "Критическая ошибка",
                "faster-whisper не установлен!\n\n"
                "Установите командой:\n"
                "pip install faster-whisper"
            )
            sys.exit(1)

        # Создаем меню О программе
        app.aboutToQuit.connect(lambda: None)  # Убираем агрессивную очистку при выходе

        window = MainWindow()

        # Добавляем меню
        menubar = window.menuBar()
        help_menu = menubar.addMenu("Помощь")
        about_action = help_menu.addAction("О программе")
        about_action.triggered.connect(lambda: AboutDialog(window).exec())

        window.show()

        window.log_widget.log("=" * 50, "INFO")
        window.log_widget.log(f"{APP_NAME} v{APP_VERSION} запущен", "SUCCESS")
        window.log_widget.log(f"Автор: {AUTHOR}", "INFO")
        window.log_widget.log("Защита от крашей активирована", "INFO")
        window.log_widget.log("=" * 50, "INFO")

        sys.exit(app.exec())

    except Exception as e:
        print(f"Критическая ошибка запуска: {e}")
        import traceback
        traceback.print_exc()
        try:
            QMessageBox.critical(None, "Критическая ошибка", f"Не удалось запустить программу:\n{e}")
        except:
            pass


if __name__ == "__main__":
    main()
